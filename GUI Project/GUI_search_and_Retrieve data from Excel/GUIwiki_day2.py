import wikipedia

#python to docx
from docx import Document
def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)
    #.summary เป็นแบบสรุปบทความข้อมูลที่จะหา **********
    data = wikipedia.summary(keyword)  

    # Page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content


    doc = Document() #สร้างไฟล์ word ใน python
    doc.add_heading(keyword,0) # 0 = ความใหญ่ของตัวอักษร heading

    doc.add_paragraph(data2) # โยน data ลงไปใน doc
    doc.save(keyword+'.docx')
    print('Generate files success')


#เปลี่ยนเป็นภาษาไทย
wikipedia.set_lang('th')
from tkinter import *
from tkinter import ttk
from tkinter import messagebox   #ไว้ try except Error

GUI = Tk()
GUI.title('โปรแกรม wiki')
GUI.geometry('400x300')
#--------Config----------
FONT1 = ('Angsana New',15)
#--------คำอธิบาย-----------
L = ttk.Label(GUI,text='ค้นหาบทความ',font=FONT1)
L.pack()
#--------ช่องค้นหาข้อมูล---------
v_search = StringVar()
E1 = ttk.Entry(GUI,textvariable=v_search,font=FONT1,width=40)
E1.pack(pady=10)
#--------Button-----------
def Search():
    keyword = v_search.get() #.get() คือ ดึงข้อมูลเข้ามา
    try: 
        # ลองค้นหาดูส่าได้ผลลัพธ์หรือไม่ หากได้ให้ผ่าน
        language = v_radio.get() # th/en/zh
        Wiki(keyword,language)
        messagebox.showinfo('บันทึกสำเร็จ','ค้นหาข้อความสำเร็จ บันทึกเรียบร้อย')  # เป็น popup ขึ้นว่าสำเร็จ
    except:
        #หากรันคำสั่งแล้วมีปัญหา แสดงข้อความแจ้งเตือน
        messagebox.showwarning('Keyword Error','กรุณากรอกคำค้นหาใหม่')
        # สั่งให้เปิดไฟล์ หรือ ไม่เปิด ใช้ radio button
        
    #print(wikipedia.search(keyword))
    #result = wikipedia.summary(keyword)
    #print(result)
    
B1 = ttk.Button(GUI,text='Search',command=Search)
B1.pack(ipadx=20,ipady=10) # ipadx ขยายปุ่มแนวนอน


# เลือกภาษา  #เราทำปุ่มเป็นแนวนอน # สั่งให้เปิดไฟล์ หรือ ไม่เปิด ใช้ radio button
F1 = Frame(GUI) #เป็นกระดาน whiteboardแล้วเอา Button เราไปแปะ
F1.pack(pady=20)

v_radio = StringVar() # ช่องเก็บข้อมูลภาษา

RB1 = ttk.Radiobutton(F1,text='ภาษาไทย',variable=v_radio,value='th')
RB2 = ttk.Radiobutton(F1,text='ภาษาอังกฤษ',variable=v_radio,value='en')
RB3 = ttk.Radiobutton(F1,text='ภาษาจีน',variable=v_radio,value='zh')
RB1.invoke() # defualt สั่งให้ค่าเริ่มต้นเป็นภาษาไทย

#เราทำปุ่มภาษาเป็นแนวนอน เรียงเป้น Matric
RB1.grid(row=0,column=0)
RB2.grid(row=0,column=1)
RB3.grid(row=0,column=2)

#เราทำปุ่มภาษาเป็นแนวตั้ง
#RB1.pack()
#RB2.pack()
#RB3.pack()





GUI.mainloop()
