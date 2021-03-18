from docx import Document
import wikipedia


def Wiki(keyword,lang='th'):
    wikipedia.set_lang(lang)
    #.summary เป็นแบบสรุปบทความข้อมูลที่จะหา **********
    data = wikipedia.summary(keyword)  

    # Page + content บทความทั้งหน้า
    data2 = wikipedia.page(keyword)
    data2 = data2.content


    doc = Document() #สร้างไฟล์ word ใน python
    doc.add_heading(keyword,0) # 0 = ความใหญ่ของตัวอักษร heading

    doc.add_paragraph(data2) # โยน data ลงไปใน doc
    doc.save(keyword+'.docx')
    print('Generate files success')

try:
    Wiki('dfbgb','en')
except:
    print('ERROR')
        

#Wiki('huawei','zh')
#Wiki('xiaomi','zh')
